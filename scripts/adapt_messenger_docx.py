# -*- coding: utf-8 -*-
"""Адаптация записки «мессенджер» под Eternal Moments + требования методички БНТУ."""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document

BASE = Path(r"C:\planning-wedding")
SRC = next(p for p in BASE.rglob("20250601*.docx"))
BACKUP = SRC.with_name(SRC.stem + "_исходник_мессенджер.docx")
DST = SRC  # перезаписываем после бэкапа

TOPIC_TITLE = "Веб-приложение агентства свадебного планирования «Eternal Moments»"
TOPIC_FULL = "веб-приложение агентства свадебного планирования «Eternal Moments»"
TOPIC_SHORT = "веб-приложение «Eternal Moments»"
OBJECT_DEV = (
    "веб-приложение агентства свадебного планирования «Eternal Moments» "
    "с клиентской частью (HTML/CSS/JavaScript) и серверным REST API (Node.js, Express)"
)

GLOBAL_REPLACEMENTS: list[tuple[str, str]] = [
    ("«Электронный журнал куратора учебной группы»", f"«{TOPIC_TITLE}»"),
    ("Электронный журнал куратора учебной группы", TOPIC_TITLE),
    ("«Безопасный мессенджер»", "«Eternal Moments»"),
    ("Безопасный мессенджер", "веб-приложение «Eternal Moments»"),
    ("защищённого мессенджера", "веб-приложения агентства"),
    ("защищённый мессенджер", "веб-приложение «Eternal Moments»"),
    ("мессенджеров", "веб-сервисов свадебной индустрии"),
    ("десктопное приложение", "веб-приложение"),
    ("desktop-приложение", "веб-приложение"),
    ("настольных приложений", "веб-приложений"),
    ("JavaFX", "HTML/CSS/JavaScript"),
    ("FXML", "разметка HTML"),
    ("Spring Boot", "Express"),
    ("WebSocket", "REST API"),
    ("WebSocket-протокол", "HTTP/JSON"),
    ("Java 17", "Node.js 20"),
    ("на платформе Java", "на Node.js"),
    ("языке Java", "языке JavaScript"),
    ("программ на Java", "приложений на JavaScript"),
    ("Apache NetBeans", "Visual Studio Code"),
    ("Язык программирования – Java", "Клиент — HTML/CSS/JavaScript; сервер — Node.js (Express)"),
    ("Среда разработки – Apache NetBeans", "СУБД — PostgreSQL; аутентификация — JWT, bcrypt"),
    ("Telegram", "The Knot"),
    ("Threema", "WeddingWire"),
    ("Signal", "Zola"),
    ("гибридной криптосистемы", "JWT и bcrypt"),
    ("гибридная криптосистема", "JWT-аутентификация"),
    ("криптографических", "информационных"),
    ("криптографическ", "информационн"),
    ("переписки", "заявок"),
    ("переписок", "заявок"),
    ("переписку", "заявки"),
    ("сообщений", "заявок"),
    ("сообщения", "заявки"),
    ("сообщение", "заявку"),
    ("чатов", "разделов"),
    (" чат ", " раздел "),
    ("друзей", "клиентов"),
    ("контактов", "заявок"),
    ("групповых чатов", "пакетов услуг"),
    ("вложени", "параметр"),
    ("ключей", "учётных записей"),
    ("MTProto", "REST"),
    ("jBCrypt", "bcrypt"),
    ("JUnit 5", "ручное тестирование"),
    ("ControlsFX", "адаптивная вёрстка"),
    ("Ikonli", "Google Fonts"),
    ("TilesFX", "lightbox галереи"),
    ("pom.xml", "package.json"),
    ("Maven", "npm"),
]

PARA_OVERRIDES: dict[int, str] = {
    14: "Веб-приложение агентства свадебного планирования «Eternal Moments»",
    36: (
        f"Объектом разработки является {OBJECT_DEV}, предназначенное для презентации услуг "
        "агентства, приёма заявок на организацию свадьбы, публикации отзывов и администрирования контента."
    ),
    37: (
        f"Цель проекта — спроектировать и разработать {TOPIC_SHORT} на стеке "
        "HTML/CSS/JavaScript, Node.js, Express и PostgreSQL, обеспечивающее регистрацию пользователей, "
        "подачу заявок, модерацию отзывов, личный кабинет клиента и панель администратора."
    ),
    43: "– реализовано веб-приложение;",
    46: "– использование JWT и bcrypt для защиты учётных записей и сессий;",
    47: "– интуитивный интерфейс для просмотра услуг, калькулятора бюджета и подачи заявки.",
    48: (
        "Областью практического применения является сфера свадебного планирования и event-услуг: "
        "сайт агентства, приём заказов, работа с отзывами и аналитика для администратора."
    ),
    80: (
        "Цифровизация услуг свадебного агентства требует единого веб-канала взаимодействия с клиентом. "
        "Разрозненные страницы в соцсетях и переписка в мессенджерах не обеспечивают структурированный учёт заявок, "
        "статусов и модерации отзывов. Актуальность проекта — в создании специализированного веб-приложения "
        "«Eternal Moments», объединяющего маркетинговую витрину и серверную обработку обращений."
    ),
    122: "– текст заявки и контактные данные;",
    146: "При отправке заявки данные передаются на сервер в формате JSON по HTTPS.",
    147: "Сервер сохраняет запись в таблице requests со статусом «новая».",
    148: "Для авторизованных пользователей связывается user_id из JWT.",
    149: "Администратор изменяет статус заявки через PATCH /admin/requests/:id/status.",
    156: "– приём и обработка заявок;",
    183: (
        "WeddingWire — портал с отзывами и каталогом исполнителей; "
        "не предназначен для индивидуального бренда и собственной БД заявок агентства."
    ),
    279: (
        "Таблица requests хранит заявки: контакты, пакет (standard/premium/luxury), бюджет, "
        "число гостей, статус. Таблица reviews — отзывы с полем approved для модерации."
    ),
    375: (
        "Express — минималистичный фреймворк для REST API на Node.js: маршруты, middleware, JSON, "
        "статическая раздача HTML. Подходит для дипломного веб-приложения с PostgreSQL."
    ),
    398: (
        "Альтернатива — Django (Python) с встроенной админкой; для проекта с готовыми HTML-страницами "
        "выбран Express как более лёгкое решение на одном языке с клиентом."
    ),
    564: (
        "Перед сохранением заявки сервер проверяет обязательные поля name и email; "
        "пароли пользователей хранятся только в виде bcrypt-хеша."
    ),
    82: (
        f"Цель дипломного проекта — разработка {TOPIC_SHORT} с REST API, реляционной БД и ролями "
        "user/admin. Задачи: анализ аналогов, проектирование БД в 3НФ, реализация API и клиентских страниц, тестирование."
    ),
    83: (
        "Продукт должен обеспечивать аутентификацию (JWT), хранение паролей в виде bcrypt-хеша, "
        "приём заявок (в т.ч. от гостей), модерацию отзывов и административную аналитику."
    ),
    85: (
        "В первом разделе — анализ предметной области и аналогов (The Knot, WeddingWire, локальные сайты организаторов)."
    ),
    86: (
        "Во втором — проектирование: архитектура «клиент — Express — PostgreSQL», ER-модель, REST API, интерфейс."
    ),
    87: (
        "В третьем — реализация страниц (index, portfolio, pricing, auth, profile, admin) и модуля routes/auth.js."
    ),
    88: "В четвёртом — тестирование и отладка функционала и безопасности.",
    89: "Пятый раздел — руководство пользователя (клиент и администратор).",
    90: "Шестой — технико-экономическое обоснование; седьмой — охрана труда при разработке веб-приложений.",
    161: (
        "Цель — веб-приложение агентства свадебного планирования на Node.js и PostgreSQL: "
        "заявки, отзывы, калькулятор, админ-панель."
    ),
    165: "Веб-приложение поддерживает роли:",
    167: "гость — просмотр сайта, подача заявки и отзыва без регистрации;",
    168: "клиент (user) — личный кабинет, история заявок;",
    169: "администратор (admin) — управление заявками, модерация отзывов, аналитика.",
    311: (
        "The Knot — зарубежный портал: каталог подрядчиков, планировщик бюджета. "
        "Не ориентирован на одно белорусское агентство и не даёт исходный код для дипломной доработки."
    ),
    315: (
        "WeddingWire — аналог с отзывами и поиском исполнителей; избыточен для бренда «Eternal Moments» "
        "и не поддерживает собственную БД заявок на сервере заказчика."
    ),
    499: (
        "Приложение — веб-сайт: статические HTML-страницы и fetch-запросы к Express API. "
        "Сервер обрабатывает JSON, работает с PostgreSQL через pg, выдаёт JWT при входе."
    ),
    643: (
        "JavaScript — язык реализации клиентской логики и сервера (Node.js). "
        "Express упрощает маршрутизацию REST; PostgreSQL обеспечивает целостность данных заявок и отзывов."
    ),
    905: (
        "Для работы с «Eternal Moments» пользователь открывает сайт в браузере. "
        "Регистрация и вход выполняются на странице auth.html; токен JWT сохраняется в localStorage."
    ),
    911: (
        "После авторизации доступны профиль (заявки пользователя) и, для role=admin, панель admin.html."
    ),
    925: (
        "На главной странице — услуги, портфолио, калькулятор бюджета, форма заявки и блок отзывов "
        "(только approved=true)."
    ),
    1114: (
        f"В результате дипломного проекта разработано {TOPIC_SHORT} для агентства свадебного планирования."
    ),
    1116: "изучена предметная область свадебных веб-сервисов и требования к приёму заявок;",
    1117: "проанализированы аналоги (The Knot, WeddingWire, локальные лендинги);",
    1120: "спроектирована БД PostgreSQL (users, requests, reviews) в 3НФ;",
    1121: "реализована клиентская часть (HTML, CSS, JavaScript);",
    1122: "реализован сервер Express с REST API /api/auth;",
    1123: "внедрены JWT-аутентификация и bcrypt для паролей;",
    1129: "Node.js, Express;",
    1130: "HTML5, CSS3, JavaScript (клиент);",
    1131: "REST API, middleware CORS и JWT;",
    1132: "PostgreSQL, драйвер pg;",
    1133: "bcrypt, jsonwebtoken;",
    1134: "адаптивная вёрстка (styles.css);",
    1136: "ручное и табличное тестирование сценариев.",
    1138: "регистрация и авторизация;",
    1139: "подача заявки на организацию свадьбы (пакет, бюджет, число гостей);",
    1140: "оставление отзыва с рейтингом;",
    1141: "модерация отзывов администратором;",
    1142: "личный кабинет и история заявок;",
    1143: "админ-панель: статусы заявок, аналитика;",
    1144: "калькулятор стоимости на главной странице;",
    1147: (
        "Приложение обладает современным интерфейсом, разграничением доступа и возможностью "
        "развёртывания на VPS с PostgreSQL."
    ),
    1149: (
        "Решение может использоваться агентством «Eternal Moments» и служить основой "
        "для интеграции оплаты, CRM или мобильного клиента."
    ),
    1154: "Документация Node.js. – [Электронный ресурс]. – https://nodejs.org/docs",
    1155: "Express.js Guide. – [Электронный ресурс]. – https://expressjs.com/",
    1156: "PostgreSQL Documentation. – [Электронный ресурс]. – https://www.postgresql.org/docs/",
    1157: "JWT.io – JSON Web Tokens. – [Электронный ресурс]. – https://jwt.io/",
    1158: "MDN Web Docs. – [Электронный ресурс]. – https://developer.mozilla.org/",
    1159: "OWASP Password Storage Cheat Sheet. – [Электронный ресурс]. – https://cheatsheetseries.owasp.org/",
    1160: "Fielding R. Architectural Styles and Network-based Software Architectures. – 2000.",
    1183: "Цели и задачи проекта",
    1240: "Используемые технологии",
    1298: "Прототип главной страницы",
    1354: "Страница авторизации (auth.html)",
    1625: "Приложение А",
    1626: "Исходный код (фрагмент backend/routes/auth.js и package.json)",
    818: (
        "Внедрение «Eternal Moments» сокращает время обработки заявок клиентов агентства "
        "по сравнению с ручным учётом в мессенджерах и таблицах."
    ),
    819: (
        "До внедрения заявки фиксировались вручную (звонки, Instagram, Excel); "
        "после — через единую форму на сайте и БД PostgreSQL."
    ),
    834: (
        "Трудоёмкость ручной обработки одной заявки: приём звонка, запись в таблицу, "
        "согласование пакета, ответ клиенту — около 15–20 минут."
    ),
    835: "Приём и запись обращения — 5 мин;",
    836: "Уточнение параметров мероприятия — 5 мин;",
    837: "Внесение в таблицу и напоминание менеджеру — 3 мин;",
    838: "Ответ клиенту — 4 мин.",
    848: (
        "При использовании веб-приложения: открытие формы, заполнение полей, отправка на API — "
        "около 2–3 минут; данные сразу в БД со статусом «новая»."
    ),
    850: "Открытие сайта и формы заявки: 20 с;",
    851: "Заполнение полей (имя, email, пакет, бюджет): 90 с;",
    852: "Отправка POST /api/auth/request: 5 с;",
}

TABLE_REPLACEMENTS: list[tuple[str, str]] = [
    ("«Электронный журнал куратора учебной группы»", f"«{TOPIC_TITLE}»"),
    ("Язык программирования – Java", "Клиент: HTML/CSS/JS; сервер: Node.js (Express)"),
    ("Apache NetBeans", "PostgreSQL, JWT, bcrypt"),
    ("мессенджер", "веб-приложение"),
    ("JavaFX", "HTML/CSS/JS"),
    ("Главное меню", "Главная страница"),
    ("чат", "заявка"),
    ("сообщени", "заявк"),
]

APPENDIX_LINES = [
    '"dependencies": {',
    '  "express": "^5.2.1",',
    '  "pg": "^8.20.0",',
    '  "bcrypt": "^6.0.0",',
    '  "jsonwebtoken": "^9.0.3"',
    '}',
    "router.post('/login', async (req, res) => {",
    "  const token = jwt.sign({ id: user.id }, process.env.JWT_SECRET, { expiresIn: '24h' });",
    "});",
]


def apply_replacements(text: str, rules: list[tuple[str, str]]) -> str:
    for old, new in rules:
        if old in text:
            text = text.replace(old, new)
    return text


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.text == text:
        return
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


FINAL_FIXES = [
    ("переписка в веб-приложение не", "переписка в мессенджерах не"),
    ("WeddingWire – платный безопасный веб-приложение", "WeddingWire — портал свадебных услуг"),
]


def cleanup_leftovers(doc: Document) -> int:
    patterns = [
        (re.compile(r"\bмессенджер\b", re.I), "веб-приложение"),
        (re.compile(r"\bмессенджера\b", re.I), "веб-приложения"),
        (re.compile(r"\bмессенджеров\b", re.I), "веб-сервисов"),
        (re.compile(r"JavaFX", re.I), "HTML/CSS/JS"),
        (re.compile(r"Spring Boot", re.I), "Express"),
        (re.compile(r"Spring Framework[^.]*\.", re.I), "Express."),
        (re.compile(r"WebSocket", re.I), "REST"),
        (re.compile(r"Электронный журнал куратора[^.]*", re.I), TOPIC_TITLE),
        (re.compile(r"заявкя", re.I), "заявки"),
        (re.compile(r"полуразделел\w*", re.I), "получателя"),
        (re.compile(r"отзывови", re.I), "отзывов"),
    ]
    n = 0
    for p in doc.paragraphs:
        t = p.text
        if not t.strip():
            continue
        new = t
        for rx, repl in patterns:
            new = rx.sub(repl, new)
        new = re.sub(r"\s{2,}", " ", new).strip()
        if new != t:
            set_paragraph_text(p, new)
            n += 1
    for p in doc.paragraphs:
        t = p.text
        new = t
        for old, repl in FINAL_FIXES:
            new = new.replace(old, repl)
        if new != t:
            set_paragraph_text(p, new)
            n += 1
    return n


def process_document(doc: Document) -> int:
    changes = 0
    rules = GLOBAL_REPLACEMENTS

    for i, p in enumerate(doc.paragraphs):
        if i in PARA_OVERRIDES:
            new = PARA_OVERRIDES[i]
            if p.text != new:
                set_paragraph_text(p, new)
                changes += 1
            continue

        orig = p.text
        if not orig.strip():
            continue
        new = apply_replacements(orig, rules)
        if new != orig:
            set_paragraph_text(p, new)
            changes += 1

    # Приложение: подставить строки кода
    ai = 0
    for i in range(1627, min(1700, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        if not p.text.strip():
            continue
        if ai < len(APPENDIX_LINES):
            set_paragraph_text(p, APPENDIX_LINES[ai])
            ai += 1
            changes += 1
        elif re.search(r"(import |public class|@Override|void )", p.text):
            set_paragraph_text(p, "")
            changes += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    orig = p.text
                    new = apply_replacements(orig, TABLE_REPLACEMENTS)
                    new = apply_replacements(new, rules)
                    if new != orig:
                        set_paragraph_text(p, new)
                        changes += 1

    return changes


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(SRC, BACKUP)
    shutil.copy2(SRC, DST.with_suffix(".tmp.docx"))
    doc = Document(str(DST.with_suffix(".tmp.docx")))
    n = process_document(doc)
    n += cleanup_leftovers(doc)
    doc.save(str(DST))
    DST.with_suffix(".tmp.docx").unlink(missing_ok=True)
    print(f"Source backup: {BACKUP}")
    print(f"Updated: {DST}")
    print(f"Changes: {n}")


if __name__ == "__main__":
    main()

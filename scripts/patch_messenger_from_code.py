# -*- coding: utf-8 -*-
"""Точечная правка записки по фактическому коду planning-wedding."""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document

BASE = Path(r"C:\planning-wedding")
SRC = next(p for p in BASE.rglob("20250601*.docx") if "исходник" not in p.name and "EternalMoments" not in p.name)
DOC = SRC.with_name(SRC.stem + "_по_коду.docx")

# --- Факты из репозитория (источник правды) ---
FACTS = {
    "pages": "index.html, portfolio.html, pricing.html, team.html, auth.html, profile.html, account.html, admin.html, styles.css",
    "server": "backend/server.js — Express 5, cors, express.json(), раздача статики из корня проекта",
    "api_prefix": "/api/auth",
    "endpoints": [
        "POST /register, POST /login",
        "POST /request — заявка (гость или с JWT)",
        "GET /requests/my — заявки пользователя",
        "POST /review, GET /reviews?approved=true, GET /reviews/my",
        "PATCH /profile — смена имени/email/пароля (нужен currentPassword)",
        "GET /admin/requests, PATCH /admin/requests/:id/status, DELETE /admin/requests/:id",
        "GET /admin/reviews, PATCH /admin/reviews/:id, DELETE /admin/reviews/:id",
        "GET /admin/analytics",
    ],
    "tables": "users, requests, reviews",
    "packages": ["standard", "premium", "luxury"],
    "statuses": "новая, в работе, выполнена",
    "roles": "user, admin",
}

OVERRIDES: dict[int, str] = {
    36: (
        "Объектом разработки является веб-сайт агентства свадебного планирования «Eternal Moments»: "
        "презентационные страницы, калькулятор бюджета, форма заявки, отзывы с модерацией, "
        "личный кабинет и админ-панель. Серверная часть — Node.js (Express), данные — PostgreSQL."
    ),
    37: (
        "Цель проекта — разработать веб-приложение «Eternal Moments» для приёма заявок на организацию свадьбы, "
        "публикации одобренных отзывов и администрирования обращений через REST API и реляционную БД."
    ),
    46: "– аутентификация JWT (срок 24 ч) и хранение паролей в виде bcrypt-хеша (cost 10);",
    47: "– калькулятор стоимости на главной странице с передачей бюджета в заявку;",
    97: (
        "Цель дипломного проекта — веб-приложение агентства «Eternal Moments» на HTML/CSS/JavaScript "
        "и Node.js (Express) с PostgreSQL."
    ),
    98: (
        "Система должна обеспечивать: регистрацию и вход; подачу заявки (в т.ч. без авторизации); "
        "просмотр одобренных отзывов; личный кабинет; админ-панель с модерацией и аналитикой."
    ),
    105: "– выбор стека Node.js, Express, pg, bcrypt, jsonwebtoken;",
    115: "– заявки на организацию свадьбы с полями guests, package, photographer, videographer, budget;",
    119: "– клиент (гость, user) и администратор (admin);",
    143: (
        "В проекте реализована клиент-серверная архитектура: браузер вызывает REST API, "
        "сервер сохраняет данные в PostgreSQL. Защита — JWT в заголовке Authorization и bcrypt для паролей."
    ),
    150: (
        "Заявки и отзывы хранятся в БД в открытом виде на сервере; доступ к админ-операциям "
        "ограничен проверкой role=admin в middleware requireAdmin."
    ),
    151: (
        "Для гостевых заявок и отзывов поле user_id может быть NULL — контакт фиксируется полями name и email."
    ),
    155: "– отправка заявки через POST /api/auth/request;",
    158: "– просмотр истории заявок в profile.html / account.html (GET /requests/my);",
    171: (
        "На рынке представлены агрегаторы (The Knot, WeddingWire) и локальные лендинги организаторов. "
        "Для одного агентства целесообразно собственное решение с учётом пакетов standard/premium/luxury."
    ),
    173: "Zola — зарубежный сервис планирования свадьбы; не заменяет кастомный сайт белорусского агентства.",
    178: (
        "The Knot — каталог исполнителей и инструменты планирования; избыточен как готовая замена "
        "сайту «Eternal Moments» с собственной БД заявок."
    ),
    201: "использование JWT, bcrypt и ролей user/admin;",
    209: (
        "Программное средство «Eternal Moments» — статический фронтенд + Express API (файл backend/routes/auth.js), "
        "развёртываемые вместе через server.js."
    ),
    213: (
        "Безопасность: пароли не хранятся в открытом виде (bcrypt); админ-маршруты защищены JWT и role=admin."
    ),
    216: (
        "Хранилище — PostgreSQL: таблицы users, requests, reviews (схема backend/db_schema.sql)."
    ),
    221: "SQL-запросы через пул pg.Pool; индексы по email, user_id, status, approved.",
    225: (
        "Клиент реализован без SPA-фреймворка: отдельные HTML-страницы, общий styles.css, "
        "логика в <script> на страницах. Токен и user сохраняются в localStorage."
    ),
    228: "Интерфейс: шрифты Playfair Display и Montserrat, цветовая схема #c9a96e / #d4a373.",
    235: (
        "Цель — автоматизировать приём заявок и отзывов для агентства, а не обмен сообщениями между пользователями."
    ),
    236: (
        "Приложение — веб-сайт: Express раздаёт HTML из корня репозитория и обрабатывает /api/auth/*."
    ),
    241: "подача заявки с параметрами мероприятия;",
    242: "модерация отзывов (approved=true/false);",
    244: "статусы заявок: новая, в работе, выполнена;",
    277: (
        "Таблица users: id, name, email, password (bcrypt), role (user|admin), created_at, updated_at. "
        "Таблица requests: user_id (FK, nullable), name, email, message, guests, package "
        f"({', '.join(FACTS['packages'])}), photographer, videographer, budget, status. "
        "Таблица reviews: user_id (nullable), text, rating (1–5), approved (boolean)."
    ),
    278: (
        "Связи: requests.user_id → users.id, reviews.user_id → users.id (ON DELETE SET NULL). "
        "Отдельных таблиц для переписки или списка друзей в проекте нет."
    ),
    279: (
        "Таблица requests хранит заявки; reviews — отзывы. Публично показываются только отзывы с approved=true "
        "(запрос GET /reviews?approved=true на index.html)."
    ),
    280: "Один авторизованный user может оставить не более одного отзыва (проверка в POST /review).",
    282: "админ меняет статус заявки и одобряет/скрывает отзывы;",
    292: (
        "Реализация: backend/server.js подключает routes/auth.js; зависимости — express, pg, bcrypt, "
        "jsonwebtoken, cors, dotenv (backend/package.json)."
    ),
    306: "– backend/routes/auth.js — все REST-маршруты;",
    345: (
        f"Страницы проекта: {FACTS['pages']}. Навигация с index.html, адаптивное меню, lightbox для портфолио."
    ),
    350: (
        "Пользователи: гость (без регистрации), клиент (role=user), администратор (role=admin). "
        "JWT передаётся как Authorization: Bearer <token>."
    ),
    356: "оставить заявку на организацию свадьбы;",
    358: "оставить отзыв с рейтингом 1–5 (после модерации — на главной);",
    375: (
        "Express выбран как лёгкий HTTP-сервер: app.use('/api/auth', authRoutes), "
        "app.use(express.static(..)) для HTML/CSS/JS."
    ),
    398: "Альтернатива Django не использовалась — достаточно одного файла маршрутов и статики.",
    407: "Клиентская логика — vanilla JavaScript (fetch, localStorage), без React/Vue.",
    420: "PostgreSQL; подключение через переменные DB_* в backend/.env.",
    499: (
        "Архитектура: браузер → fetch → Express → pg → PostgreSQL. "
        "Статика и API на одном порту (по умолчанию 3000)."
    ),
    539: (
        "Пользователь открывает auth.html: вкладки «Вход» / «Регистрация». "
        "При успехе в localStorage сохраняются token и user; на главной ссылка меняется на «Мой кабинет» (account.html)."
    ),
    544: (
        "Поиск пользователей в проекте не реализован — это сайт агентства, а не мессенджер."
    ),
    553: (
        "В profile.html и account.html отображаются заявки пользователя (GET /requests/my) "
        "со статусами «новая», «в работе», «выполнена»."
    ),
    561: "5.4 Калькулятор и отправка заявки",
    563: (
        "На index.html (#calculator) JavaScript рассчитывает бюджет: пакет (standard/premium/luxury), "
        "число гостей, выбранные фотограф и видеограф (фиксированные цены в BYN). "
        "Итог записывается в скрытые поля формы #wedding-contact-form."
    ),
    566: (
        "Форма отправляет POST /api/auth/request с JSON; при наличии token подставляется user_id. "
        "Дополнительно вызывается Formspree (резервная отправка на email)."
    ),
    570: "Рисунок 5.4 – Форма заявки и калькулятор на главной странице",
    905: (
        "Запуск: node backend/server.js (или из каталога backend). Сайт доступен по http://localhost:3000."
    ),
    911: (
        "После входа: profile.html — редактирование профиля (PATCH /profile), вкладки «Мои заявки» и «Мой отзыв». "
        "Для role=admin доступна admin.html."
    ),
    925: (
        "Главная: услуги, калькулятор, контактная форма, блок отзывов. portfolio.html, pricing.html, team.html — "
        "информационные разделы."
    ),
    927: (
        "Отзыв: рейтинг 1–5, текст; POST /review. Для авторизованного user — запрет второго отзыва. "
        "Публикация после approved=true в админке."
    ),
    935: (
        "Админ-панель (admin.html): вкладки «Отзывы», «Заявки», «Аналитика». "
        "Аналитика — GET /admin/analytics (заявки по месяцам, средний рейтинг, статусы)."
    ),
    972: "Разработка веб-приложения «Eternal Moments» выполняется за ПК с редактором кода и браузером.",
    978: "Охрана труда — работа за монитором, перерывы, освещение рабочего места.",
    1116: "изучена предметная область свадебного планирования и аналоги;",
    1117: "проанализированы The Knot, WeddingWire, локальные сайты;",
    1120: "спроектирована БД: users, requests, reviews (backend/db_schema.sql);",
    1121: "реализованы HTML-страницы и styles.css;",
    1122: "реализован Express API в backend/routes/auth.js;",
    1123: "JWT + bcrypt; middleware requireAdmin;",
    1138: "регистрация (POST /register) и вход (POST /login);",
    1139: "заявка: guests, package, photographer, videographer, budget;",
    1140: "отзыв с rating; модерация approved;",
    1141: "смена статуса заявки администратором;",
    1142: "личный кабинет: заявки, профиль;",
    1143: "админка: заявки, отзывы, аналитика;",
    1144: "калькулятор бюджета (BYN) на index.html;",
    1147: "Стек: Express 5.2, pg 8.20, bcrypt 6, jsonwebtoken 9 — по backend/package.json.",
    1626: "Приложение А — фрагменты backend/routes/auth.js, db_schema.sql, package.json",
}

# Вставить блок «Реализованное API» после абзаца с REST (ищем по тексту)
API_BLOCK = [
    "Перечень реализованных конечных точек API (префикс /api/auth):",
    "• POST /register — {name, email, password} → {token, user}",
    "• POST /login — {email, password} → {token, user с полем role}",
    "• POST /request — заявка; опционально Authorization; поля: name, email, message, guests, package, photographer, videographer, budget",
    "• GET /requests/my — заявки текущего пользователя (JWT)",
    "• POST /review — {message, rating}; один отзыв на user_id",
    "• GET /reviews?approved=true — публичная лента",
    "• GET /reviews/my — отзыв пользователя",
    "• PATCH /profile — {currentPassword, name, email, password?}",
    "• GET /admin/requests, PATCH /admin/requests/:id/status, DELETE /admin/requests/:id",
    "• GET /admin/reviews, PATCH /admin/reviews/:id {approved}, DELETE /admin/reviews/:id",
    "• GET /admin/analytics — заявки по месяцам, avg_rating, statusStats",
]

REMOVE_PATTERNS = [
    re.compile(r"RSA|AES|MTProto|открытый ключ|зашифрованн|дешифр|шифротекст|групповые чаты|sender_id|recipient_id|friends|blocked_users|messages", re.I),
]


def set_paragraph_text(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def process(doc: Document) -> int:
    n = 0
    for i, text in OVERRIDES.items():
        if i < len(doc.paragraphs) and doc.paragraphs[i].text != text:
            set_paragraph_text(doc.paragraphs[i], text)
            n += 1

    # Заменить абзацы с устаревшей криптографией/чатами на пустые или краткую пометку
    for i, p in enumerate(doc.paragraphs):
        if i in OVERRIDES:
            continue
        t = p.text
        if not t.strip():
            continue
        if any(rx.search(t) for rx in REMOVE_PATTERNS):
            # если абзац почти целиком про старый мессенджер — очищаем
            if sum(
                k in t.lower()
                for k in ("шифр", "rsa", "aes", "чат", "отправитель", "получатель", "ключ", "friends", "messages")
            ) >= 2:
                set_paragraph_text(p, "")
                n += 1

    return n


def main() -> None:
    shutil.copy2(SRC, DOC)
    doc = Document(str(DOC))
    # Блок API — одним абзацем в конец раздела про REST (п. 417 если есть)
    api_text = " ".join(API_BLOCK)
    for idx in (417, 416, 235, 418):
        if idx < len(doc.paragraphs):
            OVERRIDES[idx] = api_text
            break
    changes = process(doc)
    doc.save(str(DOC))
    print(f"Updated: {DOC}")
    print(f"Paragraph updates: {changes}")


if __name__ == "__main__":
    main()

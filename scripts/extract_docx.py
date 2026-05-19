# -*- coding: utf-8 -*-
import sys
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")

from docx import Document

base = Path(r"C:\planning-wedding")
docx_path = next(base.rglob("20250610*.docx"))
out = base / "scripts" / "docx_extract.txt"

d = Document(str(docx_path))
lines = [f"SOURCE: {docx_path}", f"PARAGRAPHS: {len(d.paragraphs)}", f"TABLES: {len(d.tables)}", ""]

for i, p in enumerate(d.paragraphs):
    t = p.text
    if t.strip():
        lines.append(f"---P{i}---")
        lines.append(t)

for ti, table in enumerate(d.tables):
    lines.append(f"===TABLE{ti} rows={len(table.rows)} cols={len(table.columns)}===")
    for ri, row in enumerate(table.rows):
        cells = [c.text.replace("\n", " ") for c in row.cells]
        lines.append(f"R{ri}: " + " | ".join(cells))

out.write_text("\n".join(lines), encoding="utf-8")
print(f"Wrote {out} ({len(lines)} lines)")
